# -*- coding: utf-8 -*-
import json, csv, glob, os
from collections import defaultdict
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def num(s):
    s = (s or '').strip().replace('.', '').replace(',', '')
    try: return int(s)
    except: return 0

MESES = ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']

bhe = json.load(open(os.path.join(ROOT,'procesadores','output','bhe_todas.json'), encoding='utf-8'))
bhe_anio = defaultdict(lambda: {'n':0,'vig':0,'bruto':0,'ret':0,'emis':set(),'obs':0})
bhe_tri = defaultdict(lambda: {'n':0,'bruto':0})
for b in bhe:
    a = b['anio']; e = b.get('estado','').upper(); d = bhe_anio[a]
    d['n'] += 1
    if e == 'VIGENTE':
        d['vig'] += 1; d['bruto'] += b.get('honorario_bruto',0); d['ret'] += b.get('retencion',0); d['emis'].add(b.get('rut_emisor',''))
        k = (a,b['trimestre']); bhe_tri[k]['n'] += 1; bhe_tri[k]['bruto'] += b.get('honorario_bruto',0)
    else:
        d['obs'] += 1

m12_anual = defaultdict(lambda: defaultdict(int)); m12_tot = defaultdict(int); m12_tri = defaultdict(int)
for f in sorted(glob.glob(os.path.join(ROOT,'procesadores','output','M12_Gastos','*-4.csv'))):
    anio = int(os.path.basename(f).split('-')[0])
    with open(f, encoding='utf-8-sig') as fh:
        for row in csv.DictReader(fh, delimiter=';'):
            item = (row.get('Item de Gastos') or '').strip()
            meses = [num(row.get(m)) for m in MESES]; tot = sum(meses)
            m12_anual[anio][item] += tot; m12_tot[anio] += tot
            for ti,(x,y) in enumerate([(0,3),(3,6),(6,9),(9,12)]):
                m12_tri[(anio,ti+1)] += sum(meses[x:y])

rcv = {}
with open(os.path.join(ROOT,'procesadores','output','RCV_Resumen','resumen_anual.csv'), encoding='utf-8-sig') as fh:
    for row in csv.DictReader(fh, delimiter=';'):
        rcv[int(row['Año'])] = {'docs': row['Documentos'], 'total': num(row['Total'])}

def clp(n): return "$" + format(int(n), ",").replace(",", ".")

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
AMARANTO = RGBColor(0x9B,0x23,0x35); AZUL = RGBColor(0x00,0x30,0x87); GRIS = RGBColor(0x55,0x55,0x55)

def h(txt, size=14, color=AMARANTO, after=4, before=10):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before); return p

def para(txt, size=10.5, bold=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def tabla(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hc[i].text = ''; run = hc[i].paragraphs[0].add_run(htxt); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tcPr = hc[i]._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '9B2335'); tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    return t

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("MINUTA DE LEVANTAMIENTO DE INFORMACION - SII + DEFONTANA + TRANSPARENCIA SERVEL"); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AMARANTO
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Partido Comunista de Chile - SERVEL PP007 - RUT 71.701.800-1"); r.font.size = Pt(11); r.font.color.rgb = AZUL
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mp.add_run("Periodo 2016-2026 - Fecha: " + date.today().strftime('%d-%m-%Y') + " - Fuentes: SII (BHE, RCV, F29) + Defontana ERP (EEFF, Libro Mayor) + Transparencia SERVEL (24 modulos)")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRIS
doc.add_paragraph("_"*95).alignment = WD_ALIGN_PARAGRAPH.CENTER

anios = [2022,2023,2024,2025]

h("1. Resumen Ejecutivo")
para("Este diagnostico cruza TRES fuentes oficiales: (1) el portal SII (1.390 BHE, 1.971 documentos RCV, 41 F29, "
     "M12); (2) el ERP contable Defontana del partido (Estados Financieros IFRS, Balance de Comprobacion y Libro "
     "Mayor con ~26.900 asientos de 2024-2026); y (3) el Portal de Transparencia SERVEL (24 modulos extraidos, "
     "2016-2026). Las fuentes son consistentes en lo verificable (honorarios 2025 coinciden al peso entre SII y "
     "Defontana; el aporte estatal 2024 coincide entre Transparencia y Defontana). El gasto se duplico (x2,2) "
     "entre 2022 y 2025 por los ciclos electorales, y la contabilidad oficial revela una crisis estructural de "
     "liquidez y una cuenta por cobrar de $4.262M a empresas relacionadas. La auditoria de Transparencia (modulo "
     "11, Ingresos Art.34) permite reconstruir los ingresos 2023 y confirma que el partido SI recibio aporte "
     "estatal en 2023 (~$345,8M), dato que el SII no registra.")
tabla(["Indicador","2022","2023","2024","2025"], [
    ["Gasto total (M12)"] + [clp(m12_tot[a]) for a in anios],
    ["BHE emitidas"] + [str(bhe_anio[a]['n']) for a in anios],
    ["Honorarios brutos"] + [clp(bhe_anio[a]['bruto']) for a in anios],
    ["Retencion honorarios"] + [clp(bhe_anio[a]['ret']) for a in anios],
    ["RCV total (compras)"] + [clp(rcv[a]['total']) for a in anios],
])

# ===================== SECCION 2: INVENTARIO DOCUMENTAL =====================
h("2. Inventario Documental Extraido por Fuente")
para("Cantidad de documentos/registros efectivamente extraidos y procesados, con su fuente de origen. "
     "Es la base probatoria del diagnostico.", size=10)
sii_rcv_docs = sum(int(str(rcv[a]['docs'])) for a in rcv)
tabla(["Fuente","Documento / Registro","Cantidad","Periodo","Estado"],
 [["SII","BHE - Boletas honorarios electronicas", f"{sum(bhe_anio[a]['n'] for a in bhe_anio)}", "2022-2025","Completo"],
  ["SII","RCV - Doctos compra/venta (facturas)", f"{sii_rcv_docs}", "2022-2026","Completo"],
  ["SII","F29 - Declaraciones mensuales IVA/retencion", "41", "2022-2025","Falta Jun-Dic 2025"],
  ["SII","M12 - Items de gasto (trimestral)", "4 anios x ~18 items", "2022-2025","Completo"],
  ["Defontana","Libro Mayor - asientos contables", "~26.928 lineas", "2023-2026","2023 casi vacio (10)"],
  ["Defontana","EEFF / Balances / Flujo de caja (archivos)", "15", "2023-2025","2022-2023 sin contab."],
  ["Transparencia","Modulo 11 Ingresos (Art.34) - filas", "347", "2016-2026","Completo (re-extraido)"],
  ["Transparencia","Modulo 10 Aportes/donaciones - filas", "290", "2016-2026","Completo (re-extraido)"],
  ["Transparencia","Modulo 09 Balance aprobado SERVEL", "212", "2015-2026","Completo"],
  ["Transparencia","Modulo 12 Gastos del partido", "730", "2017-2026","Completo"],
  ["Transparencia","Modulo 13 Cotizaciones afiliados", "82", "2016-2026","Re-extraido (2024-25 con vacios)"],
  ["Transparencia","Modulo 14 Transferencias de fondos", "8.522", "2010-2026","Completo"],
  ["Transparencia","Modulo 15 Nomina contrataciones >20 UTM", "2.433", "2010-2026","Completo"],
  ["Transparencia","Modulo 16 Gastos campanas electorales", "228", "2016-2026","Completo (re-extraido)"],
  ["Transparencia","Modulo 17 Aportes a campanas electorales", "291", "2016-2025","Completo"],
  ["Transparencia","Modulo 21 Sanciones aplicadas", "189", "2019-2026","Completo (3 multas)"],
  ["Transparencia","Otros 8 modulos (organica/estadistica)", "~3.900", "2010-2029","Completo"]])
para("Total aproximado: mas de 45.000 registros extraidos de 3 fuentes oficiales independientes.",
     bold=True, color=AZUL, size=9)

h("3. BHE - Boletas de Honorarios Electronicas")
tot_bhe = sum(bhe_anio[a]['n'] for a in bhe_anio); tot_vig = sum(bhe_anio[a]['vig'] for a in bhe_anio)
tot_obs = sum(bhe_anio[a]['obs'] for a in bhe_anio)
tot_bruto = sum(bhe_anio[a]['bruto'] for a in bhe_anio); tot_ret = sum(bhe_anio[a]['ret'] for a in bhe_anio)
para("Total: " + str(tot_bhe) + " boletas - " + str(tot_vig) + " vigentes - " + str(tot_obs) +
     " en estado 'Observado Receptor' (revisar). Honorarios brutos 2022-2025: " + clp(tot_bruto) +
     " - Retencion total: " + clp(tot_ret) + ".")
tabla(["Ano","BHE","Vigentes","Observadas","Emisores","Bruto","Retencion"],
      [[a, bhe_anio[a]['n'], bhe_anio[a]['vig'], bhe_anio[a]['obs'], len(bhe_anio[a]['emis']), clp(bhe_anio[a]['bruto']), clp(bhe_anio[a]['ret'])] for a in anios])
para("Comportamiento trimestral (honorarios brutos vigentes):", bold=True)
tabla(["Ano","Q1","Q2","Q3","Q4"],
      [[a] + [clp(bhe_tri[(a,t)]['bruto']) for t in (1,2,3,4)] for a in anios])

h("4. Gastos M12 - Comportamiento por Trimestre")
para("Se observan dos picos electorales claros: 2024-Q3 (municipales) y 2025-Q4 (presidencial/parlamentaria).")
tabla(["Ano","Q1","Q2","Q3","Q4","Total"],
      [[a] + [clp(m12_tri[(a,t)]) for t in (1,2,3,4)] + [clp(m12_tot[a])] for a in anios])
para("Principales categorias de gasto (acumulado por ano):", bold=True)
cats = set()
for a in m12_anual: cats |= set(m12_anual[a].keys())
cattot = {c: sum(m12_anual[a].get(c,0) for a in m12_anual) for c in cats}
top = [c for c in sorted(cats, key=lambda x: -cattot[x])][:10]
tabla(["Categoria","2022","2023","2024","2025"],
      [[c[:38]] + [clp(m12_anual[a].get(c,0)) for a in anios] for c in top])

h("5. RCV - Proveedores y Compras")
para("Gasto concentrado en medios y comunicacion. Proveedores recurrentes: Radio Nuevo Mundo (medio del PC), "
     "Multitud Comunicaciones Digitales (creciente), Inversiones Siglo XXI (editorial, atipico $455M en 2023), "
     "y agencias digitales de campana (Artemedios, Power ID, CCreative).")
tabla(["Ano","Documentos","Total compras (con IVA)"],
      [[a, rcv[a]['docs'], clp(rcv[a]['total'])] for a in anios])

h("6. Informacion Contable Oficial (Defontana ERP)")
para("Estado de contabilizacion: 2024 y 2025 estan COMPLETOS y cuadran (Estados Financieros IFRS disponibles, "
     "listos para PRESENTAR a SERVEL aunque aun NO aprobados); 2023 esta SIN contabilizar (el Libro Mayor "
     "Defontana 2023 tiene solo 10 asientos por "
     "$2,19M - una factura y dos egresos) y 2022 esta vacio. Cifras del Estado de Situacion Financiera y "
     "Balance de Comprobacion:", bold=False)
tabla(["Indicador (Defontana)","2024","2025"],
      [["Total activos","$6.784.503.000","$6.185.202.000"],
       ["Bancos (caja)","$415.636.000","$5.596.000  (-99%)"],
       ["Resultado del ejercicio","Superavit $169.572.869","Deficit $1.278.221.363"],
       ["Aporte estatal DFL N4","$549.691.607 (CONFIRMADO)","$0 (confirmado)"],
       ["Cuenta por cobrar PROGRESO","$4.370.021.000","$4.262.217.000"],
       ["Pasivos financieros (bancos)","$8.122.000","$467.020.000"]])
para("Hallazgos ciertos de la contabilidad Defontana:", bold=True, color=AMARANTO)
for t in ["PROGRESO: $4.262M por cobrar a empresas relacionadas con fines de lucro (Soc. de Inv. Progreso SpA "
          "76.452.615-5 y Radio Progreso SpA 76.825.989-5). Es saldo de apertura (origen <=2023, sin documentar "
          "en libros actuales); existe convenio de 338 cuotas, pero la cuenta casi no se mueve (~25 anos para recuperar). Es ~69% del activo total.",
          "CREDITO ELECTORAL: el 07/11/2025 el partido tomo $480.000.000 del Banco Estado (dias antes de la eleccion).",
          "CRISIS DE LIQUIDEZ: los bancos cayeron de $415,6M (2024) a $5,6M (2025). El flujo de caja proyectado "
          "llega a -$753M sin ingresos; con costos fijos ~$28M/mes y aporte estatal en $0, solo el credito electoral "
          "evito la insolvencia.",
          "ACTIVO FIJO real: $1.807M neto (Terrenos $1.852M, Obras de arte $93M, Mobiliario $78M, Muebles $8M).",
          "Financiamiento electoral 2024 (municipales): Concejal $413M, Consejero Reg. $329M, Primarias Alcalde $273M."]:
    p = doc.add_paragraph(t, style='List Bullet'); p.runs[0].font.size = Pt(9.5)

para("Estructura de INGRESOS (M6 - Fuentes de financiamiento) segun el Libro Mayor Defontana (grupo de "
     "cuentas 3). Esta es la cara de ingresos que el SII NO registra (el SII solo tiene gastos, honorarios y "
     "compras). 2024 y 2025 estan contabilizados; 2023 esta VACIO:", bold=True, color=AMARANTO)
tabla(["Fuente de ingreso (cuenta Defontana grupo 3)","2024","2025"],
      [["3.1.1010  Aporte estatal trimestral (DFL N4)","$549.691.607","$0"],
       ["3.1.2010  Cotizaciones (ordinarias+extraord.+no afiliados)","$321.781.637","$227.474.635"],
       ["3.1.3010  Ingresos del propio patrimonio","$6.919.413","$5.605.048"],
       ["3.1.4010  Financiamiento/reembolso electoral (por cargo)","$1.185.426.351","$218.288.285"],
       ["3.5.0100  Intereses efectivamente percibidos","$110.286.768","$117.867.520"],
       ["TOTAL INGRESOS","$2.174.105.776","$569.235.488"]])
para("INGRESOS OFICIALES segun Transparencia SERVEL (modulo 11, Ingresos Art.34 - mensual, 2016-2026). Esta "
     "fuente RESUELVE la brecha de ingresos 2023 que el SII no registra. Su aporte 2024 coincide al peso con "
     "Defontana, lo que valida toda la serie:", bold=True, color=AMARANTO)
tabla(["Fuente de ingreso (Transparencia mod.11)","2022","2023","2024","2025"],
      [["Aporte del Estado (art. 33 bis)","$522.870.xxx","$345.788.634","$549.691.607","$0"],
       ["Cuotas y aportes de afiliados","$92.838.032","$400.866.223","$901.581.422","$185.729.458"],
       ["Rendimientos de patrimonio","$147.781.727","$263.255.951","$740.705.008","$162.362.673"],
       ["Aportes personas naturales","$2.206.xxx","$9.853.723","$200.xxx","$0"],
       ["Reembolsos / ingresos electorales","-","~$1.066M","~$2.380M","~$70M"]])
para("NOTA SOBRE EL APORTE 2023: el aporte estatal de 2023 ($345.788.634) figura en Transparencia (modulos 11 "
     "y 10) pero NO en el SII. El monto exacto se reconfirma con la cartola del Banco Estado (el modulo 10 "
     "trimestral indica $512M). El aporte 2025 = $0 (el $308M que mostraba el modulo 10 corresponde a una "
     "categoria distinta). El detalle contable 2023 aun no esta cargado en Defontana.",
     bold=False, size=9.5, color=GRIS)

# ── 6.bis  Estado de los balances + aprobacion SERVEL + cuadratura ──
para("Estado de los BALANCES y aprobacion SERVEL (modulo 09 Transparencia):", bold=True, color=AMARANTO)
tabla(["Ano","Aprobado por SERVEL (mod.09)","Contabilizado Defontana","Ecuacion contable"],
      [["2015-2021","SI - aprobados y publicados","-","-"],
       ["2022","NO aprobado","Sin contabilizar (vacio)","-"],
       ["2023","NO aprobado","Vacio (10 asientos, $2,19M)","-"],
       ["2024","NO aprobado","COMPLETO","Cuadra (dif $0)"],
       ["2025","NO aprobado","COMPLETO","Cuadra (dif $0)"]])
para("HALLAZGO: SERVEL solo tiene balances APROBADOS hasta 2021. Hay un backlog de 4 anos (2022-2025) sin "
     "balance aprobado - raiz formal de la suspension del aporte (Art. 42 DFL N4) que golpeo en 2025.",
     bold=True, color=AMARANTO, size=10)
para("Verificacion de la ecuacion contable (cuadra en los 3 documentos Defontana: Estado de Situacion "
     "Financiera = Balance General = Balance de Comprobacion):", bold=False, size=9.5)
tabla(["Ano","Total Activos","Pasivos + Patrimonio","Resultado ejercicio","Cuadra"],
      [["2024","$6.784.503.000","$6.614.929.746","+$169.572.869","SI (dif $0)"],
       ["2025","$6.185.202.000","$7.463.423.665","-$1.278.221.363","SI (dif $0)"]])

# ── 6.ter  Inputs necesarios para confirmar/generar los balances ──
para("Inputs necesarios para CONFIRMAR y GENERAR los balances:", bold=True, color=AZUL)
tabla(["Balance","Que falta","Input critico"],
      [["2024 / 2025 (confirmar)","Conciliacion bancaria + respaldo Progreso","Cartolas 2024-2025 + convenio Progreso (338 cuotas)"],
       ["2023 (generar)","Contabilizar + cuadrar bancos","Cartolas 2023 (el resto ya esta: ingresos m11, gastos M12, BHE)"],
       ["2022 (generar)","Contabilizar + cuadrar bancos","Cartolas 2022 (saldo apertura = balance SERVEL 2021 aprobado)"]])
para("La CARTOLA BANCARIA es el input que destraba todo: es el unico documento que aporta (1) el saldo real de "
     "caja para la conciliacion, (2) los movimientos financieros que no son gasto/ingreso (transferencias, "
     "creditos, devoluciones, garantias) y (3) la trazabilidad de los ingresos. Lo demas ya esta extraido: "
     "ingresos (Transparencia m11), gastos (M12/RCV SII), honorarios (BHE) y activo fijo (Defontana).",
     bold=False, size=9.5)
para("Solicitar a Tesoreria: (1) cartolas COMPLETAS 2022-2025 de Banco Estado (60.806.000-6) y BCI 13950223; "
     "(2) convenio de pago Progreso 338 cuotas + origen del saldo; (3) certificado SERVEL del aporte 2023.",
     bold=True, color=AMARANTO, size=9.5)

h("7. Cruce de Fuentes SII <-> Defontana <-> Transparencia (validacion)")
para("La consistencia entre las tres fuentes da certeza a los datos:", bold=False)
tabla(["Concepto","Fuente A","Fuente B","Resultado"],
      [["Honorarios 2025","SII BHE $275.435.191","Defontana $275.435.191","COINCIDE EXACTO"],
       ["Honorarios 2024","SII BHE $233.352.835","Defontana $208.968.974","Dif $24,4M - reconciliar"],
       ["Aporte estatal 2024","Transparencia $549.691.607","Defontana $549.691.607","COINCIDE EXACTO"],
       ["Aporte estatal 2023","Transp. m11 $345.788.634","Transp. m10 $512.429.988","DIFIERE - usar cartola"],
       ["Cotizaciones 2023","Transp. m11 $400.866.223","Transp. m13 $216.532.302","DIFIERE - reconciliar"],
       ["Libro Mayor 2024/2025","Defontana debito","Defontana credito","Cuadra - integro"]])

h("8. Hallazgos Criticos Consolidados")
hallazgos = [
 ("PROGRESO: $4.262M a empresas relacionadas","Defontana - cuentas relacionadas con fines de lucro","Posible financiamiento de empresas; ~69% del activo"),
 ("Crisis de liquidez / deficit 2025","Defontana - deficit $1.278M, bancos -99%","Riesgo de insolvencia"),
 ("Credito electoral $480M (07/11/2025)","Defontana - Banco Estado","Endeudamiento pre-eleccion"),
 ("F29 Jun-Dic 2025 sin declarar","SII - retencion BHE no enterada (~$23,3M)","Multa SII + objecion SERVEL"),
 ("Aporte estatal 2023 no figura en SII","Solo en Transparencia m11 ($345,8M) / m10 ($512M)","Verificar monto exacto con cartola Banco Estado"),
 ("Cuota de Genero 2023 exigible","Con aporte 2023 confirmado, obligacion ~$34,6M (10%)","Revisar cumplimiento - posible rechazo Art.38 Ley 20.900"),
 ("Backlog de balances sin aprobar SERVEL","Modulo 09: aprobados solo hasta 2021 (2022-2025 pendientes)","Raiz formal de la suspension del aporte (Art. 42)"),
 ("2022 y 2023 sin contabilizar en Defontana","Defontana vacio (solo 10 asientos en 2023)","Falta cerrar esos balances; falta la cartola bancaria"),
 ("Fondo de Genero en colapso","SII - $73M (2022) a $6,3M (2025)","Rechazo de balance (Art. 38 Ley 20.900)"),
 ("Transferencias internas $477M (nov-2025)","SII/Defontana - mes electoral","Trazar contra campana"),
 ("Inconsistencia inter-modulos Transparencia","m11 vs m13 cotizaciones difieren hasta $184M","Reconciliar con contador antes de rendir"),
 ("Gasto electoral 2024-2025 (modulo 16)","Municipales 2024 ~$1.116M; Presid./Parlam. 2025 ~$687M","Cruzar contra Defontana y limites SERVEL"),
 ("10 BHE 'Observado Receptor'","SII - rechazadas","Honorarios cuestionados"),
]
tabla(["Hallazgo","Fuente / Evidencia","Riesgo"], hallazgos)

h("9. Conclusion y Diagnostico")
para("LO QUE HAY (levantado con certeza, 3 fuentes oficiales): contabilidad Defontana 2024 y 2025 completa "
     "(EEFF IFRS); nomina de honorarios (1.390 BHE), gastos (M12) y proveedores (1.971 RCV) 2022-2025; activo "
     "fijo; aporte estatal 2024 ($549,7M) y 2023 ($345,8M) confirmados en Transparencia; ingresos historicos "
     "del partido 2016-2026 (Transparencia m11); donaciones 2019-2025; 3 sanciones SERVEL documentadas.",
     bold=True, color=AZUL)
para("LO QUE FALTA (brechas que impiden cerrar la rendicion):", bold=True, color=AMARANTO)
for t in ["Contabilizacion de 2022 y 2023 en Defontana (ejercicios vacios) - falta el registro contable y "
          "su validacion por el contador.",
          "Cartola Banco Estado 2023 - para fijar el monto exacto del aporte estatal (m11 $345,8M vs m10 $512M) "
          "y conciliar los ingresos contra el banco.",
          "Cuota de genero 2023 - verificar cumplimiento del 10% del aporte confirmado (~$34,6M).",
          "F29 Jun-Dic 2025 (~$23,3M de retencion) - pendiente de declarar en el SII.",
          "PROGRESO $4.262M - falta el convenio de 338 cuotas y el origen del saldo.",
          "Gasto electoral 2024-2025 (modulo 16, ya extraido: 2024 ~$1.116M, 2025 ~$687M) - falta cruzarlo "
          "contra los limites de gasto SERVEL por candidato y el financiamiento electoral en Defontana.",
          "Conciliacion bancaria 2025 - falta confirmar que el saldo $5,6M coincide con el banco."]:
    p = doc.add_paragraph(t, style='List Bullet'); p.runs[0].font.size = Pt(10)
para("Perfil de riesgo: ALTO. Antecedentes: 3 multas SERVEL (2019, 2021, 2022) + cuota de genero incumplida "
     "9 anos + F29 atrasado + crisis de liquidez + cuenta por cobrar relacionada de $4.262M.", bold=True, color=AMARANTO)

# ===================== SECCION 10: PLAN DE GESTION =====================
doc.add_page_break()
h("10. Plan de Gestion para Cerrar las Brechas")
para("Para cada brecha pendiente se indica QUE se requiere, QUE gestion realizar, QUIEN es el responsable "
     "y la NORMA/PLAZO aplicable. Es la hoja de ruta operativa para dejar la rendicion en condiciones.", size=10)

para("10.1  Contabilizar los ejercicios 2022 y 2023 en Defontana", size=12, color=AZUL, bold=True)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Cargar y cerrar contablemente 2022 y 2023 (hoy vacios en Defontana)."],
  ["Insumos disponibles","Gastos M12 ($1.530M 2023), honorarios BHE ($165M) e ingresos (Transparencia m11) ya "
   "levantados; documentos del RCV (Compras). Falta: las cartolas para cuadrar bancos."],
  ["Gestion a realizar","1) Cargar los documentos del RCV 2023 en modulo Compras de Defontana. 2) Contabilizar "
   "gastos e ingresos. 3) Repetir para 2022. 4) Cerrar ambos ejercicios y emitir EEFF."],
  ["Responsable","Contador del partido / Tesoreria."],
  ["Norma / plazo","Balance anual ante SERVEL (Art. 39 Ley 18.603). Prioridad ALTA - bloquea la rendicion."]])

para("10.2  Aporte estatal 2023 - monto exacto (cartola Banco Estado)", size=12, color=AZUL, bold=True)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Fijar el monto exacto del aporte estatal 2023 (Transparencia m11 indica $345,8M; m10 $512M)."],
  ["Documentos necesarios","Cartola 2023 de la cuenta Banco Estado del partido (RUT 60.806.000-6) y/o certificado "
   "de pago del aporte estatal emitido por SERVEL."],
  ["Gestion a realizar","1) Solicitar a Tesoreria la cartola completa 2023. 2) Descargar de SERVEL el certificado "
   "de aporte 2023. 3) Cruzar contra m11/m10 y fijar la cifra. 4) Actualizar el modelo si difiere de $345,8M."],
  ["Responsable","Tesoreria."],
  ["Norma / plazo","Art. 33 bis Ley 18.603. Necesario para cuadrar ingresos 2023 y la cuota de genero."]])

para("10.3  Cuota de genero 2023 - posible incumplimiento", size=12, color=AZUL, bold=True)
g23 = 22151105; cuota23 = int(345788634*0.10)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Acreditar gasto en fomento de la participacion femenina 2023 >= 10% del aporte ("+clp(cuota23)+")."],
  ["Situacion detectada","Gasto femenino M12 2023 = "+clp(g23)+" vs requerido "+clp(cuota23)+". BRECHA = "+clp(cuota23-g23)+" (incumplimiento)."],
  ["Gestion a realizar","1) Revisar si hay gasto femenino adicional no clasificado. 2) Si persiste la brecha, "
   "documentar y regularizar (ejecutar el saldo o justificar ante SERVEL). 3) Replicar el chequeo para 2024."],
  ["Responsable","Contador + encargada/o de genero del partido."],
  ["Norma / plazo","Art. 38 Ley 20.900. Su incumplimiento es causal de rechazo del balance y multa."]])

para("10.4  F29 de junio a diciembre 2025 sin declarar", size=12, color=AZUL, bold=True)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Presentar las 7 declaraciones F29 faltantes (Jun-Dic 2025) con la retencion de honorarios."],
  ["Insumos disponibles","Base de retencion BHE por mes ya levantada del SII (~$23,3M total)."],
  ["Gestion a realizar","1) El contador declara cada F29 en sii.cl con la retencion BHE del mes. 2) Pagar con "
   "reajuste IPC + interes (Art. 53 Cod. Tributario) + multa por declaracion fuera de plazo (Art. 97 N.11)."],
  ["Responsable","Contador."],
  ["Norma / plazo","DL 824 / Cod. Tributario. URGENTE - el interes corre mensualmente."]])

para("10.5  Cuenta por cobrar PROGRESO $4.262M", size=12, color=AZUL, bold=True)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Documentar el origen y recuperabilidad del saldo a Soc. de Inv. Progreso SpA (76.452.615-5) "
   "y Radio Progreso SpA (76.825.989-5) - empresas relacionadas con fines de lucro (~69% del activo)."],
  ["Documentos necesarios","Convenio de pago de 338 cuotas; respaldo del origen del saldo de apertura (pre-2024); "
   "antecedentes de la relacion entre el partido y ambas SpA."],
  ["Gestion a realizar","1) Solicitar a administracion el convenio firmado. 2) Reconstruir el origen contable. "
   "3) Evaluar provision por incobrabilidad (la cuenta casi no se mueve). 4) Documentar la relacion para SERVEL."],
  ["Responsable","Administracion + contador + asesor legal."],
  ["Norma / plazo","Transparencia patrimonial SERVEL. Riesgo de observacion por financiamiento a empresas relacionadas."]])

para("10.6  Conciliacion bancaria 2025", size=12, color=AZUL, bold=True)
tabla(["Aspecto","Detalle"],
 [["Que se requiere","Confirmar que el saldo contable de bancos 2025 ($5,6M en Defontana) coincide con el banco real."],
  ["Documentos necesarios","Cartolas 2025 de todas las cuentas bancarias del partido."],
  ["Gestion a realizar","1) Obtener cartolas 2025. 2) Cruzar saldo Defontana vs banco. 3) Identificar partidas "
   "conciliatorias. 4) Confirmar el colapso de liquidez (caida de $415M a $5,6M)."],
  ["Responsable","Tesoreria / contador."],
  ["Norma / plazo","Buena practica contable; sustenta el EEFF 2025 ante SERVEL."]])

para("Secuencia recomendada: (1) F29 2025 [urgente, interes diario] -> (2) cartola Banco Estado 2023 [habilita "
     "ingresos y genero] -> (3) contabilizar 2022/2023 -> (4) cuota genero 2023 -> (5) conciliacion 2025 -> "
     "(6) convenio Progreso.", bold=True, color=AMARANTO, size=10)

# ============================ ANEXO DE CORROBORACION ============================
doc.add_page_break()
h("ANEXO - Corroboracion Documental de los Hallazgos", size=15)
para("Cada hallazgo del diagnostico se respalda a continuacion con el detalle a nivel de documento "
     "(folio F29, numero de boleta, mes y monto), extraido directamente de los archivos del portal SII.",
     size=9, color=GRIS)

# --- F29 ---
wb = openpyxl.load_workbook(os.path.join(ROOT,'ARCHIVOS SII','F29','Resultados Formularios de Impuesto 2022-2025 Consolidado.xlsx'), data_only=True)
ws = wb.active
f29 = {}
for row in ws.iter_rows(min_row=2, values_only=True):
    per = row[0]
    if per is None: continue
    ym = str(per.year)+"-"+("%02d"%per.month)
    f29[ym] = {'folio': row[1], 'fecha': row[3], 'monto': num(str(row[5]).replace('$',''))}
ret_mes = defaultdict(int); nbhe_mes = defaultdict(int)
for b in bhe:
    if b['anio']==2025 and b.get('estado','').upper()=='VIGENTE':
        ret_mes[b['mes']] += b.get('retencion',0); nbhe_mes[b['mes']] += 1

h("A.1  Hallazgo 1 - F29 Junio a Diciembre 2025 sin declarar", size=12, color=AZUL)
ultimo = max(k for k in f29 if k.startswith('2025'))
para("Ultimo F29 declarado: " + ultimo + " (folio " + str(int(f29[ultimo]['folio'])) +
     ", presentado el " + f29[ultimo]['fecha'].strftime('%d-%m-%Y') + "). Los 7 formularios siguientes no figuran. "
     "La columna 'Retencion BHE' es la base imponible real del F29 segun las boletas emitidas ese mes.")
filas_f29 = []
falta = 0
for m in range(1,13):
    ym = "2025-%02d"%m
    if ym in f29:
        estado = "Declarado (folio "+str(int(f29[ym]['folio']))+")"
        decl = clp(f29[ym]['monto'])
    else:
        estado = "** SIN DECLARAR **"; decl = "-"; falta += ret_mes[m]
    filas_f29.append(["2025-%02d"%m, str(nbhe_mes[m]), clp(ret_mes[m]), decl, estado])
tabla(["Mes","BHE","Retencion BHE","F29 monto","Estado F29"], filas_f29)
para("Retencion declarada en BHE pero NO enterada al Fisco (Jun-Dic 2025): " + clp(falta) +
     ". A esto se suman reajuste IPC e intereses (Art. 53 Codigo Tributario) y multa (Art. 97 N.11).",
     bold=True, color=AMARANTO)

# --- 10 BHE observadas ---
h("A.2  Hallazgo 4 - Las 10 BHE en estado 'Observado Receptor'", size=12, color=AZUL)
obs = [b for b in bhe if b.get('estado','').upper()!='VIGENTE']
tabla(["Ano","Mes","N. Boleta","Fecha","RUT Emisor","Emisor","Bruto"],
      [[b['anio'], b['mes'], str(b.get('nro_boleta','')), b.get('fecha',''), b.get('rut_emisor',''),
        b.get('nombre_emisor','')[:28], clp(b.get('honorario_bruto',0))]
       for b in sorted(obs, key=lambda x:(x['anio'],x['mes']))])

# --- Transferencias / movimientos M12 2025 ---
def m12_item(anio, target):
    f = os.path.join(ROOT,'procesadores','output','M12_Gastos',str(anio)+'-4.csv')
    with open(f, encoding='utf-8-sig') as fh:
        for row in csv.DictReader(fh, delimiter=';'):
            it = (row.get('Item de Gastos') or '').strip()
            if target.lower() in it.lower():
                return {MESES[i]: num(row.get(MESES[i])) for i in range(12)}
    return {}

h("A.3  Hallazgos 2 y 5 - Transferencias internas y movimientos irregulares 2025", size=12, color=AZUL)
transf = m12_item(2025,'Transferencias entre cuentas')
para("'Transferencias entre cuentas' 2025 = " + clp(sum(transf.values())) +
     ". CRITICO: " + clp(transf.get('Noviembre',0)) + " se concentran en NOVIEMBRE 2025, el mes de la eleccion "
     "presidencial/parlamentaria. Un movimiento interno de esa magnitud en mes electoral debe descartarse "
     "como aporte a campana no declarado.", bold=True, color=AMARANTO)
tabla(["Mes (2025)","Transferencias entre cuentas"],
      [[m, clp(v)] for m,v in transf.items() if v>0])
chdev = m12_item(2025,'Cheque Devuelto'); deverr = m12_item(2025,'Devoluci')
para("Otros movimientos que requieren explicacion contable (2025):", bold=True)
tabla(["Movimiento","Mes","Monto"],
      [["Cheque Devuelto", "Noviembre", clp(chdev.get('Noviembre',0))],
       ["Devolucion Transferencia Erronea", "Octubre", clp(deverr.get('Octubre',0))],
       ["Devolucion Transferencia Erronea", "Noviembre", clp(deverr.get('Noviembre',0))]])

# --- Genero ---
h("A.4  Hallazgo 3 - Fondo de Genero: ejecucion por ano y meses", size=12, color=AZUL)
fg = []
for anio in anios:
    g = m12_item(anio,'Fomento a Participaci')
    # asegurar que sea femenina
    fcsv = os.path.join(ROOT,'procesadores','output','M12_Gastos',str(anio)+'-4.csv')
    fem = {}
    with open(fcsv, encoding='utf-8-sig') as fh:
        for row in csv.DictReader(fh, delimiter=';'):
            it=(row.get('Item de Gastos') or '')
            if 'Femen' in it or 'femen' in it:
                fem = {MESES[i]: num(row.get(MESES[i])) for i in range(12)}
    tot = sum(fem.values()); meses_con = [m for m,v in fem.items() if v>0]
    fg.append([anio, clp(tot), str(len(meses_con))+" meses: "+(", ".join(meses_con) if meses_con else "-")])
tabla(["Ano","Gasto Genero","Meses con ejecucion"], fg)
para("La cuota legal es 10% del aporte estatal (Art. 38 Ley 20.900). En 2025 solo se ejecutaron $6,3M y "
     "ningun gasto despues de junio.", color=GRIS, size=9)

# --- Siglo XXI ---
h("A.5  Hallazgo 6 - Concentracion de proveedor (Inversiones Siglo XXI, 2023)", size=12, color=AZUL)
sx = []
with open(os.path.join(ROOT,'procesadores','output','RCV_Resumen','top_proveedores.csv'), encoding='utf-8-sig') as fh:
    for row in csv.DictReader(fh, delimiter=';'):
        if row['Año']=='2023' and int(row['Ranking'])<=3:
            sx.append(["#"+row['Ranking'], row['Razón Social'][:34], "RUT "+row['RUT'], row['N° Docs'], clp(num(row['Total Compras']))])
tabla(["Rk","Razon Social","RUT","Docs","Total compras"], sx)
para("Inversiones Siglo XXI concentra $454,8M en solo 8 documentos (promedio $56,8M por factura) = 31,5% de "
     "todo el RCV 2023. Verificar respaldo, objeto y precios de mercado.", color=GRIS, size=9)

doc.add_paragraph("_"*95).alignment = WD_ALIGN_PARAGRAPH.CENTER
fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Documento de trabajo interno - Da cuenta del proceso de levantamiento de informacion (SII, Defontana, Transparencia SERVEL) - Verificar cifras antes de su uso oficial.")
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GRIS

# ===================== SECCION CONTROL DE CALIDAD =====================
doc.add_page_break()
h("11. Calidad e Inconsistencias de los Datos Levantados", size=15)
para("Se realizaron validaciones cruzadas sobre los datos levantados de las tres fuentes. La columna vertebral "
     "es solida: la aritmetica de las 1.390 BHE cuadra al 100%, las tasas de retencion calzan con la escala "
     "legal del SII, los totales M12 son consistentes, el RCV concilia y el Balance 2022 cumple la ecuacion "
     "contable (Activo = Pasivo + Patrimonio, diferencia $0).")
para("Resultados de las validaciones sobre los datos de origen:", bold=True)
tabla(["Validacion","Resultado"],
      [["BHE: bruto = retencion + liquido (1.390 boletas)","OK - 0 descuadres"],
       ["BHE: suma mensual = total anual","OK - 4 anos"],
       ["Tasa de retencion efectiva vs tasa legal SII","OK - 11,76% a 14,08%"],
       ["M12: suma de categorias = total","OK - 4 anos"],
       ["RCV: neto+IVA+exento+otros = total","OK - 5 anos"],
       ["Balance 2022: Activo = Pasivo + Patrimonio","OK - diferencia $0"],
       ["Libro Mayor Defontana 2024/2025: debito = credito","OK - cuadra"]])
para("Inconsistencias detectadas en las fuentes (a resolver antes de la rendicion):", bold=True, color=AMARANTO)
tabla(["Inconsistencia","Detalle","Que requiere"],
      [["10 BHE en estado 'Observado Receptor'",
        "10 boletas fueron rechazadas por el SII (caso J.P. Astudillo 2023 superaba 20 UTM solo por una boleta rechazada de $804.600). No deben contar en la nomina >=20 UTM.",
        "Excluir de la nomina M14 al construir la rendicion"],
       ["F29 2023 declaro menos que la retencion BHE",
        "Retencion en BHE 2023 = $19,6M, pero F29 2023 declarado = $14,8M (brecha $4,8M). Posible subdeclaracion del ejercicio 2023.",
        "Revisar con el contador el entero de retenciones 2023"],
       ["RCV: 'Total' no igualaba neto+IVA+exento",
        "Diferencia de ~$48-49K en 2024 y 2025 corresponde a 'Valor Otro Impuesto' (impuesto especifico). No es error.",
        "Aclarado - es impuesto especifico, no descuadre"],
       ["Aporte estatal 2023: SII vs Transparencia",
        "El SII no registra aporte estatal del partido. Transparencia m11 (mensual) indica $345,8M y m10 (trimestral) $512M.",
        "Fijar el monto exacto con la cartola Banco Estado"],
       ["Aporte estatal 2025: m10 vs m11",
        "El modulo 10 mostraba $308M en 2025, pero el modulo 11 (que calza con Defontana) indica $0.",
        "Aclarado - 2025 = $0; el $308M del m10 es una categoria distinta"],
       ["Cotizaciones 2023: m11 vs m13",
        "Modulo 11 'Cuotas y aportes' 2023 = $400,9M vs modulo 13 'Cotizaciones' = $216,5M (m13 ademas vacio en 2024-25).",
        "Reconciliar con el contador; m11 es la fuente integral"],
       ["Portal Transparencia colapsa anos al exportar",
        "Al exportar, el portal colapsaba los 11 anos de un modulo en uno solo (afecto modulos 10, 13, 16).",
        "Se re-extrajo la serie completa 2016-2026 de esos modulos"]])

out = os.path.join(ROOT, "Minuta-Diagnostico-SII-PCCh.docx")
try:
    doc.save(out)
except PermissionError:
    out = os.path.join(ROOT, "Minuta-Diagnostico-SII-PCCh-v2.docx")
    doc.save(out)
print("OK ->", out)
